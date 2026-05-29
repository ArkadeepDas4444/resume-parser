from docx import Document

def extract_text_from_docx(docx_path):
    extracted_text = []

    try:
        document = Document(docx_path)

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                extracted_text.append(text)

        for table in document.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    extracted_text.append(" | ".join(row_text))

        return "\n".join(extracted_text).strip()

    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return None
